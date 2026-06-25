"""
app/lore/readers.py

File readers for the document ingestion pipeline.
Supports: .docx (python-docx), .pdf (PyMuPDF), .txt/.md (plain read).

All readers return plain text. For DOCX/PDF, paragraph/page breaks are
preserved as double newlines so downstream chunking can use them.
"""
from pathlib import Path


class UnsupportedFileTypeError(Exception):
    pass


def read_txt(path: str | Path) -> str:
    """Read a .txt or .md file as plain text (UTF-8, fallback latin-1)."""
    path = Path(path)
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return path.read_text(encoding="latin-1")


def read_docx(path: str | Path) -> str:
    """Extract all paragraph text from a .docx file, including tables."""
    import docx  # python-docx

    path = Path(path)
    document = docx.Document(str(path))

    parts = []
    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Also pull text from tables (often used for character/faction stat blocks)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            row_text = " | ".join(c for c in cells if c)
            if row_text:
                parts.append(row_text)

    return "\n\n".join(parts)


def read_pdf(path: str | Path) -> str:
    """Extract text from a PDF, page by page, separated by double newlines."""
    import fitz  # PyMuPDF

    path = Path(path)
    parts = []
    with fitz.open(str(path)) as doc:
        for page in doc:
            text = page.get_text().strip()
            if text:
                parts.append(text)

    return "\n\n".join(parts)


def read_file(path: str | Path) -> str:
    """
    Dispatch to the correct reader based on file extension.
    Raises UnsupportedFileTypeError for unknown extensions.
    """
    path = Path(path)
    ext = path.suffix.lower()

    if ext in (".txt", ".md"):
        return read_txt(path)
    elif ext == ".docx":
        return read_docx(path)
    elif ext == ".pdf":
        return read_pdf(path)
    else:
        raise UnsupportedFileTypeError(
            f"Unsupported file type '{ext}'. Supported: .txt, .md, .docx, .pdf"
        )


def chunk_text(text: str, max_chars: int = 6000, overlap: int = 200) -> list[str]:
    """
    Split long text into chunks suitable for Claude extraction calls.
    Splits on paragraph boundaries (double newline) where possible to
    avoid cutting entities/sentences mid-way. Adds small overlap between
    chunks so entities spanning a boundary aren't missed.
    """
    if len(text) <= max_chars:
        return [text]

    paragraphs = text.split("\n\n")
    chunks = []
    current = ""

    for para in paragraphs:
        if len(current) + len(para) + 2 <= max_chars:
            current = f"{current}\n\n{para}" if current else para
        else:
            if current:
                chunks.append(current)
                # Carry overlap from end of previous chunk (P2 fix: validate combined length)
                if overlap:
                    overlap_text = current[-overlap:]
                    candidate = overlap_text + "\n\n" + para
                    if len(candidate) > max_chars:
                        # Para alone is fine; skip the overlap to stay under limit
                        current = para
                    else:
                        current = candidate
                else:
                    current = para
                # If the paragraph itself exceeds max_chars, hard-split it immediately
                if len(current) > max_chars:
                    for i in range(0, len(current), max_chars):
                        chunks.append(current[i:i + max_chars])
                    current = ""
            else:
                # single paragraph longer than max_chars — hard split
                for i in range(0, len(para), max_chars):
                    chunks.append(para[i:i + max_chars])
                current = ""

    if current:
        chunks.append(current)

    return chunks
